# test_docx_reader.py

import os
from docx import Document as DocxDocument # Using the same alias as your main app
from docx.opc.exceptions import PackageNotFoundError # For catching corrupted DOCX

# !!! IMPORTANT: Set this variable to the EXACT name of your DOCX file !!!
# Make sure this file is in the same directory as this test script,
# or provide an absolute path.
DOCX_PATH = "/home/pre-team/Downloads/English_Check 2.docx"
# OR, if your file is named "English_Check 1.docx":
# DOCX_FILENAME_TO_TEST = "English_Check 1.docx"
# OR, if it's in a subdirectory called 'data':
# DOCX_FILENAME_TO_TEST = os.path.join("data", "English_Check.docx")
# OR, an absolute path:
# DOCX_FILENAME_TO_TEST = "/full/path/to/your/English_Check.docx"


def test_docx_file(path):
    print(f"--- Testing DOCX file: '{path}' ---")

    if not os.path.exists(path):
        print(f"❌ ERROR: File does not exist at the specified path: '{path}'")
        print("   Please check the DOCX_FILENAME_TO_TEST variable and ensure the file is in the correct location.")
        return

    print(f"✅ File found at: '{path}'")

    try:
        print("\nAttempting to open and read the document...")
        doc = DocxDocument(path)
        
        paragraphs_text = []
        if doc.paragraphs:
            for i, para in enumerate(doc.paragraphs):
                para_text = para.text.strip()
                if para_text: # Only consider non-empty paragraphs
                    paragraphs_text.append(para_text)
                    if i < 5: # Print first 5 non-empty paragraphs for a quick check
                        print(f"  [Para {i+1} Preview]: {para_text[:100]}...") # Print first 100 chars
        
        if not paragraphs_text:
            print("\n⚠️  Successfully opened the document, but no text content was found in its paragraphs.")
            print("   This might mean the document is empty, contains only images/tables without text in paragraphs, or the text is structured differently.")
        else:
            print(f"\n✅ Successfully read {len(paragraphs_text)} non-empty paragraphs from the document.")
            print(f"   Total characters read from paragraphs: {sum(len(p) for p in paragraphs_text)}")

        # You can add more checks here, e.g., for tables if your document has them
        # print(f"\nDocument has {len(doc.tables)} tables.")

        print("\n--- DOCX Read Test COMPLETED ---")

    except FileNotFoundError:
        # This should ideally be caught by os.path.exists, but as a fallback
        print(f"❌ ERROR: FileNotFoundError. The system could not find the file at '{path}'.")
    except PackageNotFoundError:
        print(f"❌ ERROR: 'PackageNotFoundError'. This often means the DOCX file '{path}' is corrupted or not a valid Office Open XML file.")
        print("   Try opening the file in Microsoft Word or another DOCX editor to check its integrity.")
        print("   If it opens fine, ensure it's saved as a standard .docx (not .doc or other formats).")
    except Exception as e:
        print(f"❌ ERROR: An unexpected error occurred while trying to read '{path}':")
        print(f"   Error type: {type(e).__name__}")
        print(f"   Error details: {e}")
        print("   The file might be corrupted, in an unexpected format, or there might be an issue with the python-docx library.")

if __name__ == "__main__":
    if DOCX_FILENAME_TO_TEST == "English_Check.docx" and not os.path.exists("English_Check.docx"):
         # A common case is the user had "English_Check 1.docx" in the error
        if os.path.exists("English_Check 1.docx"):
            print("ℹ️ NOTE: Default DOCX_FILENAME_TO_TEST is 'English_Check.docx' but 'English_Check 1.docx' exists.")
            print("   If your file is 'English_Check 1.docx', please update DOCX_FILENAME_TO_TEST in this script.")
            print("   Proceeding with testing 'English_Check.docx' for now...\n")
        elif DOCX_FILENAME_TO_TEST == "YOUR_DOCX_FILENAME_HERE.docx": # A placeholder if the user hasn't changed it
            print("⚠️  Please update the 'DOCX_FILENAME_TO_TEST' variable in this script to your actual .docx file name.")
            print("   Example: DOCX_FILENAME_TO_TEST = \"MyDocument.docx\"")

    test_docx_file(DOCX_FILENAME_TO_TEST)